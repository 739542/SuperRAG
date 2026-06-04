from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document


DOC_PATH = Path(r"E:\Dify\docs\_tmp_review_design.docx")
ORIGINAL_PATH = Path(r"E:\Dify\docs\提交文档\SuperRAG系统设计说明_提交版.docx")
BACKUP_PATH = Path(r"E:\Dify\docs\提交文档\SuperRAG系统设计说明_提交版.bak.docx")


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text


def set_cell_text(cell, text: str) -> None:
    cell.text = text


def set_row_texts(table, row_index: int, values: list[str]) -> None:
    row = table.rows[row_index]
    for cell, value in zip(row.cells, values, strict=True):
        set_cell_text(cell, value)


def main() -> None:
    doc = Document(DOC_PATH)

    # Version and provenance updates
    set_paragraph_text(
        doc.paragraphs[1],
        "系统名称：本地部署企业知识库问答 / SuperRAG / dify-lite 联调项目 主入口：http://127.0.0.1:8088/ 项目根目录：E:\\Dify 文档版本：V1.1",
    )
    set_row_texts(
        doc.tables[1],
        1,
        [
            "V1.1",
            "2026-06-02",
            "依据本地最新前后端代码修订接口清单、设计辅助分层、设计输出字段和文档删除状态描述。",
            "",
            "已按当前项目源码复核",
        ],
    )
    set_paragraph_text(
        doc.paragraphs[35],
        "本文档不是通用 RAG 系统模板说明，而是围绕当前项目的真实目录、文件职责、接口链路、启动方式和联调状态整理。本文已结合 E:\\Dify 本地最新前后端代码进行逐项复核；仍未完全展开的部分，会在相应章节以“待补充项”或“需继续核验”形式标识。",
    )
    set_paragraph_text(
        doc.paragraphs[62],
        "当前设计说明以 E:\\Dify 本地最新项目代码、目录结构和联调状态为依据，已对前端 service、dify-lite 路由与服务层、Docker 配置和部署脚本做逐项核对。对于仍未完全闭环或尚未形成正式规范的接口、字段和持久化行为，本文继续采用谨慎描述，并在 3.4 节集中列出需要后续补充确认的内容。",
    )

    # Process / design paragraphs
    set_paragraph_text(
        doc.paragraphs[147],
        "前后端主要通过 JSON 接口交换信息。文档上传使用 multipart/form-data；场景接口返回结构化 JSON；引用原文接口通过 query 参数传递 chunkId、documentId、sourceName。设计辅助场景当前采用“后端分步生成 + schema 校验/修复 + 前端 tab 化渲染 + Mermaid 图示渲染”的组织方式，前端会根据结构化字段呈现业务分析、功能清单、用例、模块、数据与权限、追踪关系、风险、后续动作和图示结果。",
    )
    set_paragraph_text(
        doc.paragraphs[172],
        "设计辅助当前定位为“结构化设计输出 + 浏览器端图示渲染”。后端 design_pipeline 会结合检索证据分步生成业务分析、功能清单、详细文本用例、模块建议、数据对象、权限分析、异常场景、风险、待确认问题、追踪矩阵、证据覆盖、后续动作与 diagram；随后通过 output_validator 和 app/schemas 对 JSON 做修复、schema 归一和质量评估。前端 design-service.js 与 app.js 会把这些字段映射为业务分析、功能清单、详细文本用例、模块划分建议、数据与权限、需求追踪、图示输出、风险与待确认问题、后续动作建议等 tab；其中 diagram 使用 Mermaid 文本图示方案，渲染失败时保留源码查看和复制能力。",
    )
    set_paragraph_text(
        doc.paragraphs[200],
        "本节汇总当前项目静态解析结论。内容已结合 E:\\Dify 本地最新源码进行核对；仍需二次确认或尚未形成正式规范的项目，在 3.4 节列出。",
    )
    set_paragraph_text(
        doc.paragraphs[210],
        "本文档已按正式系统设计说明书结构组织，包含封面、签署页、修改记录、目录、文档总说明、系统设计、接口设计、数据组织设计、需求追踪、项目解析报告、图片清单和待补充项。文档内容已结合当前本地最新 SuperRAG / dify-lite 代码复核，未将完整 Dify Docker 栈误写为 8088 主页面直接业务后端，未将前端 fallback/mock 误写为真实后端闭环，并已补充设计辅助 pipeline 与真实接口清单的最新状态。",
    )

    # Component / construct tables
    set_row_texts(
        doc.tables[16],
        2,
        [
            "UI-02",
            "部分真实接入 + fallback",
            "api.js、api-config.js、document/chat/design/training/handover/history/dashboard/settings service",
            "接口封装、字段映射、原文弹窗、多模式结果适配、Mermaid 渲染。",
            "/api 接口、本地兼容逻辑。",
            "需避免将 fallback 写成后端闭环；其中 design-service 已适配更完整的结构化字段。",
        ],
    )
    set_row_texts(
        doc.tables[16],
        8,
        [
            "SVC-06",
            "可用但有风险/场景服务",
            "POST /api/scenes/design、design_pipeline.py、output_validator.py、design-service.js",
            "设计辅助分步生成、JSON 修复、schema 校验、质量评估和 Mermaid 图示输出。",
            "模型服务、前端 Mermaid、schema 归一逻辑。",
            "当前已不是单一 prompt + 简单解析链路，而是第一版 pipeline 化实现。",
        ],
    )
    set_row_texts(
        doc.tables[17],
        5,
        [
            "设计服务构件",
            "design-service.js",
            "设计辅助结构化结果、多字段映射和 Mermaid 图示渲染。",
            "前端服务",
            "已适配业务分析、数据与权限、需求追踪等扩展字段，仍保留 fallback。",
        ],
    )
    set_row_texts(
        doc.tables[17],
        6,
        [
            "场景服务构件",
            "training-service.js / handover-service.js",
            "培训模式、交接模式输出。",
            "前端服务",
            "training 走真实场景接口并保留前端 fallback；handover 已接入 handover_pipeline，失败时可回退。",
        ],
    )
    set_row_texts(
        doc.tables[17],
        9,
        [
            "场景协调构件",
            "frontend_service.py",
            "run_scene、get_document_source、legacy runner 与 scene pipeline 协调。",
            "后端服务",
            "前后端契约核心；design/handover 会优先进入 pipeline。",
        ],
    )

    # UI / process updates
    set_row_texts(
        doc.tables[20],
        5,
        [
            "设计辅助区域",
            "业务分析、功能清单、用例、模块、数据与权限、需求追踪、风险、后续动作、图示 tab。",
            "可用但有风险",
            "Mermaid 渲染失败时保留源码；字段数量和质量受模型输出与证据覆盖影响。",
        ],
    )
    set_row_texts(
        doc.tables[21],
        5,
        [
            "P-05",
            "设计辅助流程",
            "前端提交设计问题；后端检索证据；design_pipeline 分步生成 JSON；output_validator 与 schema 做修复和归一；前端按多个 tab 渲染并执行 Mermaid 图示渲染。",
            "POST /api/scenes/design、design_pipeline.py、output_validator.py、design-service.js、app.js",
            "结构化设计结果、质量评估和 Mermaid 图示。",
        ],
    )

    # External interface tables
    set_row_texts(
        doc.tables[22],
        1,
        ["页面与静态资源接口", "I-EXT-01", "HTTP 页面接口", "高", "浏览器", "dify-lite routes.py"],
    )
    set_row_texts(
        doc.tables[22],
        2,
        ["系统与集合配置接口", "I-EXT-02", "HTTP API", "高", "第一版前端 / 运维检查", "dify-lite health/config/collections"],
    )
    set_row_texts(
        doc.tables[22],
        3,
        ["文档管理接口", "I-EXT-03", "HTTP API", "高", "第一版前端 document-service.js", "dify-lite documents/source/import/delete"],
    )
    set_row_texts(
        doc.tables[22],
        4,
        ["场景与检索接口", "I-EXT-04", "HTTP API", "高", "第一版前端 chat/design/training/handover service", "dify-lite scenes/retrieval/chat"],
    )
    set_row_texts(
        doc.tables[22],
        5,
        ["模型服务接口", "I-EXT-05", "HTTP API", "高", "dify-lite chat_service.py", "外部模型服务"],
    )

    set_row_texts(
        doc.tables[23],
        1,
        ["I-EXT-01", "GET /、GET /<asset_path>", "同步请求", "HTTP", "返回第一版前端页面及静态资源。", "主入口为 8088。"],
    )
    set_row_texts(
        doc.tables[23],
        2,
        ["I-EXT-02", "GET /api/health、GET /api/config、GET/POST /api/collections", "同步请求", "HTTP JSON", "提供健康检查、公开配置和知识库集合管理。", "当前真实 routes.py 已实现。"],
    )
    set_row_texts(
        doc.tables[23],
        3,
        ["I-EXT-03", "GET /api/documents、POST /api/documents/import、DELETE /api/documents/<document_id>、GET /api/documents/source", "同步请求", "HTTP / multipart/form-data", "提供文档列表、导入、删除和原文追溯。", "删除接口已存在，但 Weaviate 同步删除仍未实现。"],
    )
    set_row_texts(
        doc.tables[23],
        4,
        ["I-EXT-04", "POST /api/scenes/general|training|handover|design、POST /api/retrieval/query、POST /api/chat/completions", "同步请求", "HTTP JSON", "执行场景化问答、检索测试和模型回答。", "training/handover/design 都有真实后端入口，但部分前端仍保留 fallback。"],
    )
    set_row_texts(
        doc.tables[23],
        5,
        ["I-EXT-05", "由 .env 中模型服务配置决定", "后端调用", "HTTP 或兼容模型 API 协议", "调用模型生成回答、结构化设计结果或交接结果。", "Base URL / API Key / Model Name 必须有效。"],
    )

    # Data / scene capability tables
    set_row_texts(
        doc.tables[26],
        6,
        [
            "设置项",
            "前端 service / 后端配置待核验",
            "settings-service.js、.env",
            "页面设置和模型服务配置。",
            ".env 持久化；页面设置待核验",
        ],
    )
    set_row_texts(
        doc.tables[26],
        7,
        [
            "设计辅助结果",
            "接口响应 JSON / 前端渲染状态",
            "design-service.js、design_pipeline.py、output_validator.py、frontend_service.py",
            "业务分析、功能清单、用例、模块、数据与权限、追踪矩阵、图示与质量评估。",
            "待核验",
        ],
    )
    set_row_texts(
        doc.tables[27],
        5,
        [
            "D-05",
            "DesignResult",
            "businessObjects、businessRules、functionList、useCases、moduleSuggestions、dataObjects、permissionAnalysis、exceptionScenarios、risks、openQuestions、traceabilityMatrix、evidenceCoverage、nextActions、diagram、qualityAssessment。",
            "POST /api/scenes/design、design-service.js、design_pipeline.py",
            "设计辅助结构化输出，当前已明显超出旧版简化字段模型。",
        ],
    )
    set_row_texts(
        doc.tables[29],
        2,
        [
            "培训模式",
            "training-service.js",
            "POST /api/scenes/training、frontend_service.run_scene、retrieval_service、chat_service",
            "可用但有风险",
            "有真实后端场景接口，失败时前端会 fallback 到 mock。",
        ],
    )
    set_row_texts(
        doc.tables[29],
        3,
        [
            "交接模式",
            "handover-service.js",
            "POST /api/scenes/handover、frontend_service.py、handover_pipeline.py、output_validator.py",
            "可用但有风险",
            "已接入后端 pipeline；生成失败时可能回退到 legacy runner 或前端 fallback。",
        ],
    )
    set_row_texts(
        doc.tables[29],
        4,
        [
            "设计辅助",
            "design-service.js",
            "POST /api/scenes/design、frontend_service.py、design_pipeline.py、output_validator.py",
            "可用但有风险",
            "多字段结构化 JSON + Mermaid diagram；模型非规范输出时会做修复、校验和回退。",
        ],
    )
    set_row_texts(
        doc.tables[30],
        3,
        [
            "文档删除",
            "可用但有风险",
            "DELETE /api/documents/<document_id> 已删除 SQLite 记录和 uploads 原文。",
            "Weaviate 向量索引未同步删除，可能出现已删文档仍被召回的风险。",
        ],
    )
    set_row_texts(
        doc.tables[30],
        10,
        [
            "设计辅助结构化输出",
            "可用但有风险",
            "POST /api/scenes/design、design_pipeline.py、output_validator.py、design-service.js。",
            "模型返回不规范时会依赖 JSON 修复、schema 归一和 fallback；前端已适配多 tab 和 Mermaid 图示。",
        ],
    )
    set_row_texts(
        doc.tables[30],
        11,
        [
            "培训模式",
            "可用但有风险",
            "training-service.js 调用 POST /api/scenes/training，后端走 retrieval_service + chat_service。",
            "真实接口已存在，但前端仍保留 mock/fallback 路径。",
        ],
    )
    set_row_texts(
        doc.tables[30],
        12,
        [
            "交接模式",
            "可用但有风险",
            "handover-service.js、POST /api/scenes/handover、handover_pipeline.py、output_validator.py。",
            "真实 pipeline 已存在，但仍需继续核验结果质量与边界情况。",
        ],
    )

    # Tracking tables
    set_row_texts(
        doc.tables[33],
        3,
        ["F-03", "文档删除", "UI-02 / DATA-01", "DELETE /api/documents/<document_id>", "Document、Chunk", "无", "2.3.10"],
    )
    set_row_texts(
        doc.tables[33],
        8,
        ["F-08", "设计辅助", "SVC-06 / UI-02", "POST /api/scenes/design", "DesignResult", "DIFY_LITE_MODEL_*", "2.3.9"],
    )
    set_row_texts(
        doc.tables[33],
        9,
        ["F-09", "培训模式", "UI-02 / SVC-04", "POST /api/scenes/training", "SceneResult", "DIFY_LITE_MODEL_*", "2.3.8"],
    )
    set_row_texts(
        doc.tables[33],
        10,
        ["F-10", "交接模式", "UI-02 / handover pipeline", "POST /api/scenes/handover", "HandoverOutput/SceneResult", "DIFY_LITE_MODEL_*", "2.3.8"],
    )

    # Gap / appendix tables
    set_row_texts(
        doc.tables[36],
        1,
        [
            "设计辅助 pipeline 的详细时序图与 schema 展开仍不充分",
            "设计辅助最新实现理解。",
            "正文已补充 pipeline、validator 和 schema 总体说明。",
            "后续补充 design_pipeline.py、output_validator.py、app/schemas 的详细时序图与字段对照表。",
        ],
    )
    set_row_texts(
        doc.tables[38],
        5,
        [
            "构件清单",
            "index.html、app.js、styles.css、api-config.js、api.js、各前端 service、routes.py、frontend_service.py、chat_service.py、retrieval_service.py、ingestion_service.py、design_pipeline.py、handover_pipeline.py、output_validator.py、app/schemas、repository.py、db.py、document_loader.py、text_utils.py、weaviate_store.py、config.py、run.py、.env、deploy 脚本、docker-compose.yaml。",
        ],
    )
    set_row_texts(
        doc.tables[38],
        6,
        [
            "接口清单",
            "GET /、GET /<asset_path>、GET /api/health、GET /api/config、GET/POST /api/collections、GET /api/documents、POST /api/documents/import、DELETE /api/documents/<document_id>、GET /api/documents/source、POST /api/scenes/*、POST /api/retrieval/query、POST /api/chat/completions；文档详情、标签更新、reindex、历史、设置、日志等接口仍需继续核验。",
        ],
    )
    set_row_texts(
        doc.tables[38],
        10,
        [
            "未确认项",
            "SQLite 表结构细节、历史/设置/Dashboard 持久化、文档删除向量库同步、运行日志路径、认证权限设计、生产部署要求。",
        ],
    )
    set_row_texts(
        doc.tables[40],
        1,
        ["1", "设计辅助 pipeline 的详细时序图与字段 schema 对照。", "2.3.9", "中", "design_pipeline.py、output_validator.py、app/schemas、前端 design-service.js。"],
    )
    set_row_texts(
        doc.tables[40],
        3,
        ["3", "接口自动化测试结果与状态码样例。", "2.3.6", "中", "routes.py、前端 service 调用点、接口联调测试记录。"],
    )

    # Backup and save
    shutil.copy2(ORIGINAL_PATH, BACKUP_PATH)
    doc.save(DOC_PATH)
    shutil.copy2(DOC_PATH, ORIGINAL_PATH)
    print(f"Updated: {ORIGINAL_PATH}")
    print(f"Backup:  {BACKUP_PATH}")


if __name__ == "__main__":
    main()
