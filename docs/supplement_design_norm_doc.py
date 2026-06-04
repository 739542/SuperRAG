from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


DOC_PATH = Path(r"E:\Dify\docs\_design_norm.docx")


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag.endswith("}pPr"):
            continue
        p.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    clear_paragraph(paragraph)
    paragraph.add_run(text)


def insert_paragraph_after(paragraph: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def update_table_17(doc: Document) -> None:
    table = doc.tables[16]
    table.rows[1].cells[5].text = (
        r"stdout/stderr 写入 data\deploy-logs\dify-lite.stdout.log 与 "
        r"data\deploy-logs\dify-lite.stderr.log；启动器 PID 写入 "
        r"data\deploy-logs\dify-lite-launcher.pid。"
    )
    table.rows[2].cells[5].text = "导入 warnings 随接口 JSON 返回；当前未见独立导入日志表或日志文件规范。"
    table.rows[3].cells[5].text = "场景接口异常主要通过 warning / error JSON 返回；当前未见统一错误码表。"
    table.rows[4].cells[5].text = "成功时返回原文和 chunk 信息，参数缺失或无法定位文档时返回 400 JSON。"
    table.rows[5].cells[5].text = "可通过 docker compose logs 或容器日志查看；sandbox 自带 /health 健康检查。"


def update_table_18(doc: Document) -> None:
    table = doc.tables[17]
    table.rows[2].cells[3].text = "系统 Python 或 py -3 启动的 Flask 单进程服务。"
    table.rows[2].cells[4].text = "deploy.cmd 启动；当前适合本地演示与轻量联调。"
    table.rows[3].cells[4].text = "当前未见自动备份/恢复脚本，备份依赖手工复制 data 目录。"
    table.rows[4].cells[3].text = "Docker Desktop / Docker Compose；3000/5001/8080 为主要对外端口。"


def update_table_22(doc: Document) -> None:
    table = doc.tables[21]
    table.rows[2].cells[5].text = "健康检查/配置/集合查询成功 200；创建集合成功 201；name 缺失返回 400。"
    table.rows[3].cells[5].text = (
        "列表/原文查看成功 200，导入成功 201，删除成功 200；删除不存在文档返回 404；"
        "当前无 GET /api/documents/{id}。"
    )
    table.rows[4].cells[5].text = (
        "scene、retrieval、chat 正常返回 200；query 缺失或 collection_id 无效返回 400；"
        "不支持的 scene 返回 404。"
    )
    table.rows[5].cells[5].text = "模型超时或生成异常由 chat_service 封装为 warning / error，当前未见独立 error_code 枚举。"
    table.rows[6].cells[5].text = "deploy-all.ps1 会显式等待 3000/5001/8080 端口；sandbox unhealthy 不必然阻断 8088 主链路。"


def update_table_25(doc: Document) -> None:
    table = doc.tables[24]
    table.rows[4].cells[1].text = "接口响应 JSON / localStorage 摘要"
    table.rows[4].cells[2].text = "retrieval_service.py、chat_service.py、frontend_service.py、api-config.js"
    table.rows[4].cells[3].text = "回答证据链、原文追溯，以及历史摘要中的引用记录。"
    table.rows[4].cells[4].text = "后端默认不持久化；可间接进入浏览器 localStorage"

    table.rows[5].cells[1].text = "localStorage(superrag_real_history) + mock 合并"
    table.rows[5].cells[2].text = "api-config.js、history-service.js"
    table.rows[5].cells[3].text = "历史记录展示、会话筛选、场景统计。"
    table.rows[5].cells[4].text = "仅浏览器侧持久化，非服务端持久化"

    table.rows[6].cells[1].text = "后端 /api/config + 前端本地状态 / .env"
    table.rows[6].cells[2].text = "settings-service.js、api-config.js、.env"
    table.rows[6].cells[3].text = "页面设置展示、模型配置读取、Workflow 测试日志。"
    table.rows[6].cells[4].text = "后端 .env 持久化；页面编辑状态非服务端持久化"

    table.rows[7].cells[4].text = "否，当前以接口响应和 localStorage 历史摘要形式保留"


def update_table_26(doc: Document) -> None:
    table = doc.tables[25]
    rows = [
        (
            "D-01",
            "Collection",
            "id、name、description、created_at。",
            "routes.py、repository.py、ingestion_service.py",
            "SQLite collections 表的真实对象。",
        ),
        (
            "D-02",
            "Document",
            "id、collection_id、filename、original_name、title、doc_type、project、version、scene、summary、status、content_type、char_count、chunk_count、created_at。",
            "repository.py、ingestion_service.py、document-service.js",
            "SQLite documents 表的真实字段集合。",
        ),
        (
            "D-03",
            "Chunk",
            "id、document_id、collection_id、position、content、cleaned_content、token_count、metadata_json。",
            "repository.py、retrieval_service.py、text_utils.py",
            "SQLite chunks 表的真实字段集合；metadata_json 中含 source_name/file_path/chunk 参数。",
        ),
        (
            "D-04",
            "Citation",
            "chunkId、documentId、sourceName、snippet/content、relevanceScore、page。",
            "retrieval_service.py、chat_service.py、chat-service.js、training/handover/design service",
            "当前前端多模式统一消费的证据对象。",
        ),
        (
            "D-05",
            "SceneResult",
            "summary、evidence、risks、nextActions、citations、evidenceLevel、pipelineVersion、pipelineSteps、queryDesigner、retriever、validator。",
            "frontend_service.run_scene、chat_service.py、chat-service.js",
            "通用问答、培训、交接、设计四类模式的公共响应基底。",
        ),
        (
            "D-06",
            "DesignResult",
            "businessObjects、businessRules、functionList、useCases、moduleSuggestions、dataObjects、permissionAnalysis、exceptionScenarios、risks、openQuestions、traceabilityMatrix、evidenceCoverage、nextActions、diagram、qualityAssessment。",
            "design_pipeline.py、output_validator.py、design_schema.py、design-service.js",
            "设计辅助结构化输出，已明显超出旧版简化字段模型。",
        ),
        (
            "D-07",
            "LocalHistoryRecord",
            "id、sessionId、sceneMode、project、summary、originalQuestion、outputSummary、citations、createdAt。",
            "api-config.js、history-service.js、dashboard-service.js",
            "当前真实历史记录对象主要存于浏览器 localStorage 键 superrag_real_history。",
        ),
    ]
    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = value


def update_table_27(doc: Document) -> None:
    table = doc.tables[26]
    while len(table.rows) < 10:
        table.add_row()
    rows = [
        ("C-01", "DIFY_LITE_MODEL_BASE_URL", r"<SuperRAG安装目录>\dify-lite\.env", "指定模型服务基础地址。", "模型调用关键配置。"),
        ("C-02", "DIFY_LITE_MODEL_API_KEY", r"<SuperRAG安装目录>\dify-lite\.env", "指定模型服务认证 Key。", "不得写入前端。"),
        ("C-03", "DIFY_LITE_MODEL_NAME", r"<SuperRAG安装目录>\dify-lite\.env", "指定模型名称。", "需与模型服务兼容。"),
        ("C-04", "DIFY_LITE_VECTOR_STORE", r"<SuperRAG安装目录>\dify-lite\.env", "指定 local 或 weaviate 检索模式。", "决定是否访问 Weaviate。"),
        ("C-05", "DIFY_LITE_CHUNK_SIZE / CHUNK_OVERLAP", r"<SuperRAG安装目录>\dify-lite\.env", "控制切片长度与重叠。", "导入接口也可临时覆盖。"),
        ("C-06", "DIFY_LITE_MAX_CONTEXT_CHUNKS", r"<SuperRAG安装目录>\dify-lite\.env", "控制回答阶段可注入的最大证据块数。", "影响回答长度与成本。"),
        ("C-07", "前端 API 配置", r"<SuperRAG安装目录>\第一版\api-config.js", "统一前端接口地址和兼容设置。", "默认同源 /api，file:// 下退回 8088。"),
        ("C-08", "Docker Compose 配置", r"<SuperRAG安装目录>\dify\docker\docker-compose.yaml", "完整 Dify 栈容器编排。", "外围依赖。"),
        ("C-09", "Docker 环境配置", r"<SuperRAG安装目录>\dify\docker\.env", "完整 Dify 栈环境变量。", "不等同于 dify-lite .env。"),
        ("C-10", "sandbox 配置", r"<SuperRAG安装目录>\dify\docker\volumes\sandbox\conf\config.yaml", "sandbox 容器配置。", "sandbox unhealthy 不一定影响 8088 主链路。"),
    ]
    for row_idx, row_data in enumerate(rows, start=1):
        if row_idx >= len(table.rows):
            table.add_row()
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = value


def update_table_28(doc: Document) -> None:
    table = doc.tables[27]
    if len(table.rows) < 8:
        table.add_row()
    table.rows[2].cells[4].text = "后端接口真实存在，但前端失败时会回退到 mock；结果结构主要由通用 scene runner 映射。"
    table.rows[3].cells[4].text = "已接入后端 pipeline；生成失败时可能回退到 legacy runner 或前端 fallback。"
    table.rows[4].cells[4].text = "多字段结构化 JSON + Mermaid diagram；模型非规范输出时会做修复、校验和回退。"
    table.rows[6].cells[3].text = "部分可用"
    table.rows[6].cells[4].text = "documentCount/categoryCount 真实读取 /health 与 /documents；今日问答数/设计数依赖本地 history 统计。"
    table.rows[7].cells[3].text = "部分可用"
    table.rows[7].cells[4].text = "仅 /api/config 为只读真实后端；saveSettings/testWorkflow/updateWorkflow 主要是前端本地状态。"
    if len(table.rows) >= 9:
        target = table.rows[8]
    else:
        target = table.add_row()
    target.cells[0].text = "历史记录"
    target.cells[1].text = "history-service.js"
    target.cells[2].text = "window.SuperRagBackend localStorage + mock 合并"
    target.cells[3].text = "部分可用"
    target.cells[4].text = "当前未见后端历史表或历史接口，主要依赖浏览器 localStorage(superrag_real_history)。"


def update_table_29(doc: Document) -> None:
    table = doc.tables[28]
    table.rows[13].cells[1].text = "部分可用"
    table.rows[13].cells[2].text = "api-config.js 中 SuperRagBackend 将 scene 结果写入 localStorage(superrag_real_history)，history-service.js 与 mock 合并读取。"
    table.rows[13].cells[3].text = "更换浏览器或清理缓存会丢失，非服务端记录。"
    table.rows[14].cells[1].text = "部分可用"
    table.rows[14].cells[2].text = "settings-service.js 真实读取 /api/config；saveSettings/testWorkflow/updateWorkflow 仍是本地状态。"
    table.rows[14].cells[3].text = "页面设置不会直接回写 .env，也未形成后端持久化接口。"
    table.rows[15].cells[1].text = "部分可用"
    table.rows[15].cells[2].text = "dashboard-service.js 真实读取 /api/health、/api/documents，并结合 local history 计算统计。"
    table.rows[15].cells[3].text = "统计口径混合后端与前端本地数据，多用户或跨设备下不一致。"


def update_table_30(doc: Document) -> None:
    table = doc.tables[29]
    table.rows[2].cells[1].text = "部分后端接口当前确实不存在。"
    table.rows[2].cells[2].text = "文档详情、标签更新、reindex、历史持久化、设置保存、专用运行日志。"
    table.rows[2].cells[3].text = "前端通过列表派生详情、localStorage、本地状态或 mock 维持交互。"
    table.rows[2].cells[4].text = "若继续联调，应优先补齐 GET /api/documents/{id}、reindex、历史/设置持久化。"

    table.rows[6].cells[1].text = "当前缺少正式认证、审计和性能容量结论。"
    table.rows[6].cells[2].text = "生产部署、安全合规、多人协作。"
    table.rows[6].cells[3].text = "仅依赖本地目录隔离、后端 .env 保管模型凭据、deploy 脚本日志。"
    table.rows[6].cells[4].text = "正式部署前补充登录鉴权、RBAC、审计日志、备份恢复与性能压测。"


def update_table_35(doc: Document) -> None:
    table = doc.tables[34]
    rows = [
        ("正式需求规格说明编号尚未提供", "需求追踪只能使用内部 F 编号。", "本文已用 F-01 至 F-15 追踪当前实现。", "后续如有正式 SRS，可直接把 F 编号映射到正式需求编号。"),
        ("后端认证、权限、审计、脱敏和备份策略未实现", "正式部署安全性与合规性。", "正文已明确说明登录页仅为前端演示，后端未见认证/审计模块。", "后续单独补充认证模块、安全设计文档和备份恢复方案。"),
        ("文档删除与 Weaviate 同步删除未实现", "检索一致性。", "正文已明确 delete 只删除 SQLite 与 uploads，weaviate_store.py 无删除方法。", "补充向量删除接口或按 document_id / chunk_id 做批量清理。"),
        ("自动化测试与性能压测报告未发现", "验收可信度和生产评估。", "本文已改为给出真实状态码样例和最小人工回归建议。", "后续补充 pytest / 接口回归脚本 / 压测记录。"),
        ("服务端历史、设置、Dashboard 闭环未形成", "管理功能完整性。", "本文已明确 localStorage、本地状态和混合统计的真实边界。", "后续按表结构和 API 契约补齐服务端持久化。"),
    ]
    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = value


def update_table_37(doc: Document) -> None:
    table = doc.tables[36]
    table.rows[7].cells[1].text = (
        "Document、Chunk、Collection、Citation、SceneResult、DesignResult、LocalHistoryRecord。"
        "SQLite 真实表结构为 collections/documents/chunks，history 当前主要是浏览器 localStorage。"
    )
    table.rows[9].cells[1].text = (
        "正式需求编号、后端认证审计、Weaviate 删除同步、自动化测试/性能报告仍未在源码中实现或交付。"
    )


def update_table_39(doc: Document) -> None:
    table = doc.tables[38]
    headers = ["序号", "原待补充项", "补充结果", "当前结论", "依据文件"]
    for idx, value in enumerate(headers):
        table.rows[0].cells[idx].text = value
    rows = [
        ("1", "设计辅助 pipeline 的详细时序图与字段 schema 对照。", "已补充正文 2.3.9 与数据对象定义。", "当前实现为 split pipeline + validator + schema 归一 + 质量评估 + Mermaid 图示渲染。", "design_pipeline.py、output_validator.py、design_schema.py、design-service.js"),
        ("2", "SQLite 数据库表结构、字段和索引。", "已补充 2.3.7 数据组织说明。", "当前仅 3 张业务表，只有主键/唯一自动索引，无业务二级索引。", "db.py、repository.py、data/dify_lite.db"),
        ("3", "接口自动化测试结果与状态码样例。", "已补充当前真实状态码样例和验收建议。", "当前未发现自动化测试源码；文档按 routes.py 给出 200/201/204/400/404 行为说明。", "routes.py、deploy.ps1、deploy-all.ps1"),
        ("4", "文档删除是否同步删除 Weaviate 向量索引。", "已补充当前实现结论。", "delete 仅删除 SQLite 与 uploads；weaviate_store.py 未实现删除。", "routes.py、repository.py、weaviate_store.py"),
        ("5", "GET /api/documents/{id}、标签更新、reindex、运行日志接口是否存在。", "已补充接口存在性说明。", "当前不存在文档详情/tag/reindex/专用日志接口；前端使用本地状态或派生逻辑。", "routes.py、document-service.js、history-service.js、settings-service.js"),
        ("6", "历史记录、设置页、Dashboard 的后端持久化闭环。", "已补充当前闭环状态。", "history 依赖 localStorage；settings 仅只读 /api/config；dashboard 为 /health+/documents+local history 混合统计。", "api-config.js、history-service.js、settings-service.js、dashboard-service.js"),
        ("7", "用户认证、权限、日志审计和数据脱敏策略。", "已补充当前安全边界说明。", "后端未实现登录鉴权/RBAC/审计/脱敏；登录页仅为前端演示。", "index.html、app.js、routes.py、README.md"),
        ("8", "运行日志文件路径、错误码和异常处理规范。", "已补充当前日志与异常处理说明。", "deploy 脚本日志路径已明确；应用级统一错误码规范尚未形成。", "deploy.ps1、deploy-all.ps1、run.py、routes.py"),
        ("9", "Weaviate 健康检查、召回质量评估和降级策略细节。", "已补充当前降级链路说明。", "deploy-all.ps1 等待 8080；retrieval_service 捕获向量异常并回退词法检索，返回 warning。", "deploy-all.ps1、retrieval_service.py、weaviate_store.py"),
        ("10", "前端 mock/fallback 清单和正式接口替换计划。", "已补充多处 service 真实状态。", "document detail/reindex/tags、training/handover/design、history/settings/dashboard 均存在本地或 mock 路径。", "document-service.js、training-service.js、handover-service.js、design-service.js、api.js"),
        ("11", "正式测试说明、验收用例和功能覆盖报告。", "已补充当前测试现状说明。", "在第一版和 dify-lite 范围内未发现自动化测试文件，当前更接近人工联调验收。", "本地源码扫描、routes.py、前端 service 调用点"),
        ("12", "生产部署环境和并发容量要求。", "已补充当前架构边界说明。", "当前为单进程 Flask + SQLite + 本地文件，适合本地演示/轻量团队，不构成生产容量证明。", "run.py、deploy.ps1、docker-compose.yaml、db.py"),
    ]
    while len(table.rows) < len(rows) + 1:
        table.add_row()
    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = value


def main() -> None:
    doc = Document(DOC_PATH)
    paragraphs = doc.paragraphs

    set_paragraph_text(
        paragraphs[40],
        (
            r"当前设计说明以 E:\Dify 本地最新项目代码、目录结构和联调状态为依据，已对前端 service、"
            r"dify-lite 路由与服务层、Docker 配置和部署脚本做逐项核对。对于代码中确实尚未实现的认证、"
            r"审计、历史/设置持久化、Weaviate 删除同步等能力，本文不再仅以“待补充项”悬置，而是直接写出"
            r"当前真实状态、影响边界和后续建议，并在 3.4 节集中汇总本轮补充结论。"
        ),
    )

    set_paragraph_text(
        paragraphs[66],
        (
            r"当前登录页仅为前端演示入口，后端 routes.py 未提供会话登录、令牌校验或 RBAC 鉴权接口；"
            r"history/settings/dashboard 也未发现服务端用户隔离。现阶段可确认的保密性措施主要是：模型 Key "
            r"仅保存在后端 .env；上传原文和 SQLite 默认落在本地 data 目录；原文查看需通过 chunkId、"
            r"documentId、sourceName 反查，而不是直接暴露任意文件系统路径。若模型服务部署在远端，检索上下文会"
            r"被发送给外部模型 API，因此正式使用前仍需补充脱敏、权限、审计和备份策略。"
        ),
    )

    after = insert_paragraph_after(
        paragraphs[131],
        (
            r"从 routes.py 可直接归纳的接口状态码样例如下：GET /、GET /api/health、GET /api/config、"
            r"GET /api/documents、GET /api/documents/source、POST /api/scenes/*、POST /api/retrieval/query、"
            r"POST /api/chat/completions 正常返回 200；POST /api/collections 和 POST /api/documents/import 成功返回 201；"
            r"OPTIONS 预检返回 204；缺少 name、query、collection_id 等必要参数时统一返回 400；不支持的 scene 或删除不存在文档时返回 404。"
        ),
    )
    insert_paragraph_after(
        after,
        (
            r"当前 routes.py 中未实现 GET /api/documents/{id}、PATCH /api/documents/{id}/tags、"
            r"POST /api/documents/{id}/reindex、后端历史记录接口、设置保存接口和专用运行日志接口。"
            r"前端 document/history/settings/dashboard service 因此分别采用列表派生详情、localStorage、"
            r"本地状态或 mock/fallback 方式维持页面交互。"
        ),
    )

    after = insert_paragraph_after(
        paragraphs[143],
        (
            r"当前 SQLite 实际仅建立 3 张业务表：collections、documents、chunks。collections 存储知识库集合名称与描述；"
            r"documents 存储 collection_id、filename、original_name、title、doc_type、project、version、scene、"
            r"summary、status、content_type、char_count、chunk_count、created_at；chunks 存储 document_id、"
            r"collection_id、position、content、cleaned_content、token_count、metadata_json。"
        ),
    )
    after = insert_paragraph_after(
        after,
        (
            r"数据库层当前未定义除主键/唯一约束自动索引以外的业务二级索引；collections.name 依赖 UNIQUE 自动索引，"
            r"documents 与 chunks 仅依赖主键自动索引。对于本地演示和小规模文档集合这足够轻量，但在文档量上升或并发增强时，"
            r"collection_id、document_id、created_at 等字段可能需要显式索引优化。"
        ),
    )
    insert_paragraph_after(
        after,
        (
            r"前端历史记录实际保存在 localStorage 键 superrag_real_history；settings-service 中 workflows、logs、"
            r"editable retrieval/model settings 保存在前端本地状态；dashboard 则混合 /api/health、/api/documents 与本地"
            r"history 计算统计，因此这些数据当前不进入 SQLite。"
        ),
    )

    insert_paragraph_after(
        paragraphs[146],
        (
            r"通用问答、培训、交接、设计四类模式的后端场景入口已经统一到 /api/scenes/*；其中 general 直接走 "
            r"retrieval_service + chat_service，handover/design 先走 pipeline 与 validator，training 仍使用 "
            r"general scene runner 的结构化结果映射。历史页、设置页和 Dashboard 不属于同一闭环：历史依赖 localStorage，"
            r"设置仅通过 /api/config 只读读取后端配置，Dashboard 真实读取 /health 和 /documents 但问题数/设计数依赖本地 history 统计。"
        ),
    )

    after = insert_paragraph_after(
        paragraphs[150],
        (
            r"design_pipeline 的真实步骤是：1）_retrieve_scene_evidence 对 query 进行多路检索并分组；"
            r"2）_build_design_fallback 生成证据驱动的保底结构；3）_run_split_generation 并发执行 "
            r"business_analysis、use_cases、architecture_risks、traceability_quality 四个子任务；"
            r"4）safe_parse_json 对各步 JSON 做提取与修复；5）_merge_step_payloads 合并结果；"
            r"6）validate_design_output 调用 DesignOutput schema 归一字段；7）assess_quality 输出 "
            r"evidenceBoundItems、openIssueCount、score、level、canEnterReview；8）前端按 tab 渲染，并用 diagram 字段执行 Mermaid 图示渲染。"
        ),
    )
    insert_paragraph_after(
        after,
        (
            r"因此当前设计辅助并不是“单次 prompt 直接出图”的简化链路，而是“检索证据分组 + 分步生成 + "
            r"JSON 修复 + schema 归一 + 质量评估 + 前端图示渲染”的组合流程。diagram 只是最终结果中的一个字段，"
            r"失败时前端仍保留源码查看与复制能力。"
        ),
    )

    insert_paragraph_after(
        paragraphs[156],
        (
            r"当前安全、日志、mock 替换和部署容量的补充结论如下：deploy.ps1 / deploy-all.ps1 会把 dify-lite 标准输出写入 "
            r"data/deploy-logs/dify-lite.stdout.log 与 dify-lite.stderr.log，并用 pid file 管理启动器；应用级错误码尚未形成统一枚举，"
            r"routes.py 主要返回 {error: message} 的 400/404 JSON。生产级认证审计、结构化日志和容量测试在源码中仍未实现，"
            r"因此本文将其明确界定为当前版本边界，而不是文档漏项。"
        ),
    )

    insert_paragraph_after(
        paragraphs[159],
        (
            r"截至本次复核，在 E:\Dify\dify-lite 和 E:\Dify\第一版 范围内未发现 pytest、conftest、"
            r"前端 spec 或独立接口自动化用例文件。当前更接近“源码已实现 + 人工联调验收”的状态，而不是"
            r"“已有正式自动化测试报告”。因此本章追踪中的验收依据以真实路由、service 调用点和部署脚本为主，并建议至少保留"
            r"健康检查、上传、问答、引用原文、删除、场景输出六类最小回归用例。"
        ),
    )

    set_paragraph_text(
        paragraphs[177],
        r"本节汇总当前项目静态解析结论。内容已结合 <SuperRAG安装目录> 本地最新源码进行核对；3.4 节不再只是列出待补充项，而是汇总本轮已补充的结论、当前真实状态和仍需谨慎理解的边界。",
    )
    set_paragraph_text(paragraphs[183], "3.4 待补充项处理结果")
    set_paragraph_text(paragraphs[184], "表 41 待补充项处理结果")

    update_table_17(doc)
    update_table_18(doc)
    update_table_22(doc)
    update_table_25(doc)
    update_table_26(doc)
    update_table_27(doc)
    update_table_28(doc)
    update_table_29(doc)
    update_table_30(doc)
    update_table_35(doc)
    update_table_37(doc)
    update_table_39(doc)

    doc.save(DOC_PATH)


if __name__ == "__main__":
    main()
