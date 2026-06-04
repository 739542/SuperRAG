from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document


DOC_DIR = Path(r"E:\Dify\docs\提交文档")
SOURCE_PATH = DOC_DIR / "SuperRAG软件用户手册_提交版.docx"
BACKUP_PATH = DOC_DIR / "SuperRAG软件用户手册_提交版_backup.docx"
REVISED_PATH = DOC_DIR / "SuperRAG软件用户手册_修订版.docx"


def set_cell_text(cell, text: str) -> None:
    cell.text = text


def set_row_texts(table, row_index: int, values: list[str]) -> None:
    row = table.rows[row_index]
    for cell, value in zip(row.cells, values, strict=True):
        set_cell_text(cell, value)


def append_row(table, values: list[str]) -> None:
    row = table.add_row()
    for cell, value in zip(row.cells, values, strict=True):
        set_cell_text(cell, value)


def main() -> None:
    if not SOURCE_PATH.exists():
        raise FileNotFoundError(f"Source manual not found: {SOURCE_PATH}")

    shutil.copy2(SOURCE_PATH, BACKUP_PATH)
    shutil.copy2(SOURCE_PATH, REVISED_PATH)

    doc = Document(str(REVISED_PATH))

    # Table 2: 文件修改记录
    set_row_texts(
        doc.tables[1],
        2,
        [
            "A.1",
            "代码核验后修订",
            "补充真实配置路径、Docker 配置文件、未闭环接口边界、sandbox 影响范围和修订说明要求。",
            "项目组",
            "2026年6月2日",
            "形成修订版",
        ],
    )

    # Table 6: 软件清单
    append_row(
        doc.tables[5],
        [
            r"E:\Dify\dify\docker\.env",
            "完整 Dify Docker 栈的环境变量文件，供 Web、API、worker、Weaviate、sandbox 等容器读取。",
        ],
    )
    append_row(
        doc.tables[5],
        [
            r"E:\Dify\dify\docker\volumes\sandbox\conf\config.yaml",
            "sandbox 容器配置文件。主要影响完整 Dify 代码执行能力，不一定直接影响 8088 SuperRAG 主链路。",
        ],
    )

    # Table 23: 数据备份
    set_row_texts(
        doc.tables[22],
        3,
        [
            "模型配置",
            r"E:\Dify\dify-lite\.env",
            "保存模型服务地址、密钥和模型名称等关键运行配置。",
            "修改前后备份。",
        ],
    )

    # Table 26: 快速参考指南
    set_row_texts(
        doc.tables[25],
        5,
        [
            "配置模型",
            r"编辑 E:\Dify\dify-lite\.env 中 DIFY_LITE_MODEL_BASE_URL / DIFY_LITE_MODEL_API_KEY / DIFY_LITE_MODEL_NAME。",
        ],
    )

    # Table 28: 当前能力边界说明
    append_row(
        doc.tables[27],
        [
            "文档详情/标签/重新入库接口",
            "前端 document-service.js 明确标注后端当前缺少 GET /api/documents/{id}、PATCH /api/documents/{id}/tags、POST /api/documents/{id}/reindex 的完整闭环实现。",
            "正式验收时不要把这些页面交互误写成已完成后端能力，应结合 routes.py、repository.py 和后端日志再次核验。",
        ],
    )
    append_row(
        doc.tables[27],
        [
            "历史/设置/Dashboard/运行日志",
            "history-service.js、settings-service.js、dashboard-service.js 仍包含本地记录或 fallback 兼容；当前资料未显示正式运行日志接口。",
            "使用这些页面时应以 8088 主链路、文档管理页、SQLite 数据和后端实际响应为准，不应默认视为已完成持久化闭环。",
        ],
    )

    # Table 29: 待补充项
    append_row(
        doc.tables[28],
        [
            "8",
            "GET /api/documents/{id}、PATCH /api/documents/{id}/tags、POST /api/documents/{id}/reindex 的正式后端闭环状态。",
            "文档详情、标签维护、重新入库操作说明。",
            "document-service.js、routes.py、repository.py、后端联调记录。",
        ],
    )
    append_row(
        doc.tables[28],
        [
            "9",
            "历史记录、设置页、Dashboard 与运行日志接口的真实持久化和数据来源。",
            "历史、设置、概览与问题排查章节。",
            "history-service.js、settings-service.js、dashboard-service.js、routes.py、SQLite 实际运行结果。",
        ],
    )
    append_row(
        doc.tables[28],
        [
            "10",
            "8088 页面真实运行截图与原文查看弹窗截图。",
            "界面说明、验收展示和培训材料。",
            "在本地成功启动 dify-lite 后，通过浏览器访问 http://127.0.0.1:8088/ 截图替换示意图。",
        ],
    )

    doc.save(str(REVISED_PATH))
    print(f"Backup created: {BACKUP_PATH}")
    print(f"Revised manual: {REVISED_PATH}")


if __name__ == "__main__":
    main()
