from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DOC_PATH = Path(r"E:\Dify\docs\_layout_work.docx")
ORIGINAL_PATH = Path(r"E:\Dify\docs\提交文档\SuperRAG系统设计说明_提交版.docx")
BACKUP_PATH = Path(r"E:\Dify\docs\提交文档\SuperRAG系统设计说明_提交版.layout.bak.docx")


BODY_FONT = "宋体"
HEADING_FONT = "Microsoft YaHei"
COVER_FONT = "Microsoft YaHei"
ASCII_FONT = "Times New Roman"
ACCENT = RGBColor(31, 78, 121)
ACCENT_LIGHT = "D9EAF7"
TEXT_COLOR = RGBColor(44, 62, 80)


def set_run_font(run, east_asia: str, ascii_font: str, size: Pt | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def style_paragraph(paragraph, *, align=None, first_indent=None, left_indent=None, before=None, after=None, line=None, keep_with_next=None, keep_together=None) -> None:
    fmt = paragraph.paragraph_format
    if align is not None:
        paragraph.alignment = align
    if first_indent is not None:
        fmt.first_line_indent = first_indent
    if left_indent is not None:
        fmt.left_indent = left_indent
    if before is not None:
        fmt.space_before = before
    if after is not None:
        fmt.space_after = after
    if line is not None:
        fmt.line_spacing = line
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if keep_with_next is not None:
        fmt.keep_with_next = keep_with_next
    if keep_together is not None:
        fmt.keep_together = keep_together


def set_style_fonts(doc: Document) -> None:
    style_map = {
        "Normal": (BODY_FONT, ASCII_FONT, Pt(12), False, RGBColor(34, 34, 34)),
        "Heading 1": (HEADING_FONT, HEADING_FONT, Pt(18), True, ACCENT),
        "Heading 2": (HEADING_FONT, HEADING_FONT, Pt(15), True, ACCENT),
        "Heading 3": (HEADING_FONT, HEADING_FONT, Pt(13), True, TEXT_COLOR),
        "Heading 4": (HEADING_FONT, HEADING_FONT, Pt(11), True, TEXT_COLOR),
        "List Number": (BODY_FONT, ASCII_FONT, Pt(12), False, RGBColor(34, 34, 34)),
    }
    for style_name, (east_asia, ascii_font, size, bold, color) in style_map.items():
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = ascii_font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
        style.font.size = size
        style.font.bold = bold
        style.font.color.rgb = color
        fmt = style.paragraph_format
        if style_name == "Normal":
            fmt.line_spacing = 1.45
            fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(8)
        elif style_name == "Heading 1":
            fmt.space_before = Pt(18)
            fmt.space_after = Pt(10)
            fmt.keep_with_next = True
        elif style_name == "Heading 2":
            fmt.space_before = Pt(16)
            fmt.space_after = Pt(8)
            fmt.keep_with_next = True
        elif style_name == "Heading 3":
            fmt.space_before = Pt(12)
            fmt.space_after = Pt(6)
            fmt.keep_with_next = True
        elif style_name == "Heading 4":
            fmt.space_before = Pt(8)
            fmt.space_after = Pt(4)
            fmt.keep_with_next = True


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tcMar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    tbl_header = trPr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        trPr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def shade_cell(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def normalize_body_paragraphs(doc: Document) -> None:
    toc_zone = set(range(12, 31))
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        style_name = paragraph.style.name

        if idx == 0:
            style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Pt(0), before=Pt(72), after=Pt(18), line=1.25)
            for run in paragraph.runs:
                set_run_font(run, COVER_FONT, COVER_FONT, Pt(24), True, ACCENT)
            continue

        if idx == 1:
            style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Pt(0), before=Pt(0), after=Pt(28), line=1.3)
            for run in paragraph.runs:
                set_run_font(run, BODY_FONT, ASCII_FONT, Pt(11), False, TEXT_COLOR)
            continue

        if not text:
            continue

        if style_name.startswith("Heading"):
            style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=Pt(0), keep_with_next=True, keep_together=True)
            for run in paragraph.runs:
                if style_name == "Heading 1":
                    set_run_font(run, HEADING_FONT, HEADING_FONT, Pt(18), True, ACCENT)
                elif style_name == "Heading 2":
                    set_run_font(run, HEADING_FONT, HEADING_FONT, Pt(15), True, ACCENT)
                elif style_name == "Heading 3":
                    set_run_font(run, HEADING_FONT, HEADING_FONT, Pt(13), True, TEXT_COLOR)
                else:
                    set_run_font(run, HEADING_FONT, HEADING_FONT, Pt(11), True, TEXT_COLOR)
            continue

        if text.startswith("表 ") or text.startswith("图 "):
            style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Pt(0), before=Pt(8), after=Pt(6), line=1.15, keep_with_next=True)
            for run in paragraph.runs:
                set_run_font(run, HEADING_FONT, HEADING_FONT, Pt(10.5), True, ACCENT)
            continue

        if text.startswith("注：") or text.startswith("说明："):
            style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=Pt(0), before=Pt(4), after=Pt(6), line=1.3)
            for run in paragraph.runs:
                set_run_font(run, BODY_FONT, ASCII_FONT, Pt(10.5), False, RGBColor(90, 90, 90))
            continue

        if idx in toc_zone:
            style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=Pt(0), before=Pt(0), after=Pt(2), line=1.2)
            for run in paragraph.runs:
                set_run_font(run, BODY_FONT, ASCII_FONT, Pt(11), False, RGBColor(50, 50, 50))
            continue

        if style_name == "List Number":
            style_paragraph(paragraph, first_indent=Pt(0), before=Pt(0), after=Pt(4), line=1.3)
            for run in paragraph.runs:
                set_run_font(run, BODY_FONT, ASCII_FONT, Pt(12), False, RGBColor(34, 34, 34))
            continue

        style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=Pt(24), before=Pt(0), after=Pt(8), line=1.45)
        for run in paragraph.runs:
            set_run_font(run, BODY_FONT, ASCII_FONT, Pt(12), False, RGBColor(34, 34, 34))


def normalize_tables(doc: Document) -> None:
    for ti, table in enumerate(doc.tables):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        try:
            table.style = "Table Grid"
        except Exception:
            pass

        if table.rows:
            set_repeat_table_header(table.rows[0])

        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell)
                if ri == 0:
                    shade_cell(cell, ACCENT_LIGHT)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.15
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    text = paragraph.text.strip()
                    if ri == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif len(text) <= 14 and "\n" not in text and (ci == 0 or len(row.cells) <= 3):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    for run in paragraph.runs:
                        if ri == 0:
                            set_run_font(run, HEADING_FONT, HEADING_FONT, Pt(9.5), True, TEXT_COLOR)
                        else:
                            set_run_font(run, BODY_FONT, ASCII_FONT, Pt(9.5), False, RGBColor(34, 34, 34))

        if ti in {0, 1}:  # cover-page sign/revision tables breathe a bit more
            for row in table.rows:
                for cell in row.cells:
                    set_cell_margins(cell, top=130, start=140, bottom=130, end=140)


def normalize_sections(doc: Document) -> None:
    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(2.3)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.2)


def refine_cover_spacing(doc: Document) -> None:
    # Keep signature and revision pages clean and less cramped.
    for idx in [3, 8, 12]:
        if idx < len(doc.paragraphs):
            p = doc.paragraphs[idx]
            style_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=Pt(0), before=Pt(0), after=Pt(8))
            for run in p.runs:
                set_run_font(run, HEADING_FONT, HEADING_FONT, Pt(16), True, ACCENT)


def normalize_toc_page(doc: Document) -> None:
    toc_heading = doc.paragraphs[12]
    style_paragraph(toc_heading, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Pt(0), before=Pt(6), after=Pt(12), line=1.1, keep_with_next=True)
    for run in toc_heading.runs:
        set_run_font(run, HEADING_FONT, HEADING_FONT, Pt(17), True, ACCENT)

    top_level_indexes = {13, 20, 25}
    for idx in range(13, 30):
        paragraph = doc.paragraphs[idx]
        raw_text = paragraph.text.strip()
        if not raw_text:
            continue
        paragraph.text = raw_text
        if idx in top_level_indexes:
            style_paragraph(
                paragraph,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_indent=Pt(0),
                left_indent=Pt(0),
                before=Pt(3),
                after=Pt(3),
                line=1.18,
                keep_with_next=True,
            )
            for run in paragraph.runs:
                set_run_font(run, HEADING_FONT, HEADING_FONT, Pt(11.5), True, TEXT_COLOR)
        else:
            style_paragraph(
                paragraph,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_indent=Pt(0),
                left_indent=Cm(0.9),
                before=Pt(0),
                after=Pt(2),
                line=1.15,
            )
            for run in paragraph.runs:
                set_run_font(run, BODY_FONT, ASCII_FONT, Pt(10.5), False, RGBColor(68, 68, 68))

    note = doc.paragraphs[30]
    style_paragraph(note, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=Pt(0), left_indent=Pt(0), before=Pt(10), after=Pt(6), line=1.2)
    for run in note.runs:
        set_run_font(run, BODY_FONT, ASCII_FONT, Pt(9.5), False, RGBColor(110, 110, 110))


def main() -> None:
    doc = Document(DOC_PATH)
    normalize_sections(doc)
    set_style_fonts(doc)
    refine_cover_spacing(doc)
    normalize_body_paragraphs(doc)
    normalize_toc_page(doc)
    normalize_tables(doc)

    shutil.copy2(ORIGINAL_PATH, BACKUP_PATH)
    doc.save(DOC_PATH)
    shutil.copy2(DOC_PATH, ORIGINAL_PATH)
    print(f"Beautified: {ORIGINAL_PATH}")
    print(f"Backup:     {BACKUP_PATH}")


if __name__ == "__main__":
    main()
