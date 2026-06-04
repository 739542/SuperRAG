from __future__ import annotations

from pathlib import Path

from docx import Document


DOC_DIR = Path(r"E:\Dify\docs\提交文档")
SOURCE_PATH = DOC_DIR / "SuperRAG软件用户手册_修订版.docx"
TARGET_PATH = DOC_DIR / "SuperRAG软件用户手册_路径规范版.docx"
PLACEHOLDER = "<SuperRAG安装目录>"
DEFAULT_DEMO_TOKEN = "__DEFAULT_DEMO_E_DIFY__"


SPECIAL_REPLACEMENTS = {
    r"1. 将项目目录放置在 E:\Dify，确认包含“第一版”“dify-lite”“dify\docker”等目录以及部署脚本。": (
        f"1. 将项目目录放置在 {PLACEHOLDER}，确认包含“第一版”“dify-lite”“dify\\docker”等目录以及部署脚本。"
        f"{PLACEHOLDER} 表示 SuperRAG 软件在用户本机或服务器上的实际部署目录。"
        f"若采用项目组默认演示环境，该目录可为 {DEFAULT_DEMO_TOKEN}；若部署到其他路径，应以实际安装路径为准。"
    ),
    r"3. 若仅使用当前 SuperRAG 主页面，优先执行 E:\Dify\deploy.cmd 或 deploy.ps1。": (
        f"3. 若仅使用当前 SuperRAG 主页面，优先执行 {PLACEHOLDER}\\deploy.cmd 或 {PLACEHOLDER}\\deploy.ps1。"
    ),
    r"4. 若需要完整 Dify Docker 栈、Weaviate 等基础依赖，先确认 Docker Desktop 正常运行，再执行 E:\Dify\deploy-all.cmd 或 deploy-all.ps1。": (
        f"4. 若需要完整 Dify Docker 栈、Weaviate 等基础依赖，先确认 Docker Desktop 正常运行，再执行 {PLACEHOLDER}\\deploy-all.cmd 或 {PLACEHOLDER}\\deploy-all.ps1。"
    ),
    r"2. 备份 E:\Dify\dify-lite\data\dify_lite.db、E:\Dify\dify-lite\data\uploads 和 E:\Dify\dify-lite\.env。": (
        f"2. 备份 {PLACEHOLDER}\\dify-lite\\data\\dify_lite.db、{PLACEHOLDER}\\dify-lite\\data\\uploads 和 {PLACEHOLDER}\\dify-lite\\.env。"
    ),
    r"2. 进入 E:\Dify，执行 deploy-all.cmd 或 deploy-all.ps1。": (
        f"2. 进入 {PLACEHOLDER}，执行 {PLACEHOLDER}\\deploy-all.cmd 或 {PLACEHOLDER}\\deploy-all.ps1。"
    ),
    r"4. 若通过 deploy-all 启动，执行 E:\Dify\stop-all.cmd 或 stop-all.ps1 停止相关服务。": (
        f"4. 若通过 deploy-all 启动，执行 {PLACEHOLDER}\\stop-all.cmd 或 {PLACEHOLDER}\\stop-all.ps1 停止相关服务。"
    ),
    r"E:\Dify\deploy.cmd / deploy.ps1": f"{PLACEHOLDER}\\deploy.cmd / {PLACEHOLDER}\\deploy.ps1",
    r"E:\Dify\deploy-all.cmd / deploy-all.ps1": f"{PLACEHOLDER}\\deploy-all.cmd / {PLACEHOLDER}\\deploy-all.ps1",
    r"E:\Dify\stop-all.cmd / stop-all.ps1": f"{PLACEHOLDER}\\stop-all.cmd / {PLACEHOLDER}\\stop-all.ps1",
    r"E:\Dify\deploy.cmd 或 deploy.ps1": f"{PLACEHOLDER}\\deploy.cmd 或 {PLACEHOLDER}\\deploy.ps1",
    r"E:\Dify\deploy-all.cmd 或 deploy-all.ps1": f"{PLACEHOLDER}\\deploy-all.cmd 或 {PLACEHOLDER}\\deploy-all.ps1",
    r"E:\Dify\stop-all.cmd 或 stop-all.ps1": f"{PLACEHOLDER}\\stop-all.cmd 或 {PLACEHOLDER}\\stop-all.ps1",
    r"E:\Dify\dify\docker\docker-compose.yaml / .env": (
        f"{PLACEHOLDER}\\dify\\docker\\docker-compose.yaml / {PLACEHOLDER}\\dify\\docker\\.env"
    ),
    r"执行 E:\Dify\deploy.cmd 或 deploy.ps1。": (
        f"执行 {PLACEHOLDER}\\deploy.cmd 或 {PLACEHOLDER}\\deploy.ps1。"
    ),
    r"确认 Docker 正常后执行 E:\Dify\deploy-all.cmd 或 deploy-all.ps1。": (
        f"确认 Docker 正常后执行 {PLACEHOLDER}\\deploy-all.cmd 或 {PLACEHOLDER}\\deploy-all.ps1。"
    ),
    r"执行 E:\Dify\stop-all.cmd 或 stop-all.ps1；或停止对应后端进程。": (
        f"执行 {PLACEHOLDER}\\stop-all.cmd 或 {PLACEHOLDER}\\stop-all.ps1；或停止对应后端进程。"
    ),
}

GENERIC_REPLACEMENTS = [
    (r"E:\Dify\dify-lite\data\dify_lite.db", f"{PLACEHOLDER}\\dify-lite\\data\\dify_lite.db"),
    (r"E:\Dify\dify-lite\data\uploads", f"{PLACEHOLDER}\\dify-lite\\data\\uploads"),
    (r"E:\Dify\第一版", f"{PLACEHOLDER}\\第一版"),
    (r"E:\Dify\dify-lite", f"{PLACEHOLDER}\\dify-lite"),
    (r"E:\Dify\dify\docker", f"{PLACEHOLDER}\\dify\\docker"),
    (r"E:\Dify\deploy-all.cmd", f"{PLACEHOLDER}\\deploy-all.cmd"),
    (r"E:\Dify\deploy-all.ps1", f"{PLACEHOLDER}\\deploy-all.ps1"),
    (r"E:\Dify\deploy.cmd", f"{PLACEHOLDER}\\deploy.cmd"),
    (r"E:\Dify\deploy.ps1", f"{PLACEHOLDER}\\deploy.ps1"),
    (r"E:\Dify\stop-all.cmd", f"{PLACEHOLDER}\\stop-all.cmd"),
    (r"E:\Dify\stop-all.ps1", f"{PLACEHOLDER}\\stop-all.ps1"),
    (r"E:\Dify", PLACEHOLDER),
]


def normalize_text(text: str) -> str:
    updated = text
    for source, target in SPECIAL_REPLACEMENTS.items():
        updated = updated.replace(source, target)
    for source, target in GENERIC_REPLACEMENTS:
        updated = updated.replace(source, target)
    updated = updated.replace(DEFAULT_DEMO_TOKEN, r"E:\Dify")
    return updated


def walk_table(table) -> list:
    paragraphs = []
    for row in table.rows:
        for cell in row.cells:
            paragraphs.extend(cell.paragraphs)
            for nested in cell.tables:
                paragraphs.extend(walk_table(nested))
    return paragraphs


def all_paragraphs(doc: Document) -> list:
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        paragraphs.extend(walk_table(table))
    for section in doc.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
    return paragraphs


def main() -> None:
    if not SOURCE_PATH.exists():
        raise FileNotFoundError(f"Source manual not found: {SOURCE_PATH}")

    doc = Document(str(SOURCE_PATH))

    for paragraph in all_paragraphs(doc):
        original = paragraph.text
        updated = normalize_text(original)
        if updated != original:
            paragraph.text = updated

    doc.save(str(TARGET_PATH))
    print(f"Created: {TARGET_PATH}")


if __name__ == "__main__":
    main()
