from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document


def short(text: str, limit: int = 120) -> str:
    text = " ".join(text.split())
    return text[:limit]


def inspect_docx(path: Path, max_paragraphs: int, max_rows: int) -> dict:
    doc = Document(str(path))
    data: dict[str, object] = {
        "path": str(path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "paragraphs": [],
        "tables": [],
    }

    for idx, para in enumerate(doc.paragraphs[:max_paragraphs], 1):
        text = para.text.strip()
        if not text:
            continue
        data["paragraphs"].append(
            {
                "index": idx,
                "style": para.style.name if para.style else "",
                "text": short(text),
            }
        )

    for tidx, table in enumerate(doc.tables, 1):
        rows = []
        for ridx, row in enumerate(table.rows[:max_rows], 1):
            rows.append(
                {
                    "index": ridx,
                    "cells": [short(cell.text) for cell in row.cells],
                }
            )
        data["tables"].append(
            {
                "index": tidx,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "preview_rows": rows,
            }
        )

    return data


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx_path")
    parser.add_argument("--max-paragraphs", type=int, default=160)
    parser.add_argument("--max-rows", type=int, default=8)
    args = parser.parse_args()

    path = Path(args.docx_path)
    print(
        json.dumps(
            inspect_docx(path, args.max_paragraphs, args.max_rows),
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
